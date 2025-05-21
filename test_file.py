from docx import Document
import os

# Load the original document
input_path = "C:/Users/Admin/Documents/python/PGE_C2M_Unit Test Design_SOX_VEE_USG_V0.6 3.docx"

doc = Document(input_path)

# Create a new document for output
new_doc = Document()

# Copy all paragraphs first, modifying only the Execution and Actual Results section
copying = True
within_execution = False

for para in doc.paragraphs:
    text = para.text.strip()

    if "Execution and Actual Results" in text:
        within_execution = True
        copying = False
        new_doc.add_paragraph("4 Execution and Actual Results", style='Heading 1')
        new_doc.add_paragraph("4.1 Usage Rules", style='Heading 2')

        # Add corrected sequence for Usage Rules
        new_doc.add_paragraph("4.1.1 UR-CHK-MISS-INT-KWH-CONS", style='Heading 3')
        new_doc.add_paragraph("4.1.1.1 Field Value Change")
        new_doc.add_paragraph("4.1.1.1.1 BO Algorithm")
        new_doc.add_paragraph("4.1.1.1.2 Before")
        new_doc.add_paragraph("4.1.1.1.3 After")
        new_doc.add_paragraph("4.1.1.1.4 DB Results")
        new_doc.add_paragraph("4.1.1.2 List Value Change")
        new_doc.add_paragraph("4.1.1.2.1 BO Algorithm")
        new_doc.add_paragraph("4.1.1.2.2 Before")
        new_doc.add_paragraph("4.1.1.2.3 After")
        new_doc.add_paragraph("4.1.1.2.4 DB Results")

        new_doc.add_paragraph("4.1.2 UR-EXE-BASE-ALIGN-DT", style='Heading 3')
        new_doc.add_paragraph("4.1.2.1 Field Value Change")
        new_doc.add_paragraph("4.1.2.1.1 BO Algorithm")
        new_doc.add_paragraph("4.1.2.1.2 Before")
        new_doc.add_paragraph("4.1.2.1.3 After")
        new_doc.add_paragraph("4.1.2.1.4 DB Results")
        new_doc.add_paragraph("4.1.2.2 List Value Change")
        new_doc.add_paragraph("4.1.2.2.1 BO Algorithm")
        new_doc.add_paragraph("4.1.2.2.2 Before")
        new_doc.add_paragraph("4.1.2.2.3 After")
        new_doc.add_paragraph("4.1.2.2.4 DB Results")

        new_doc.add_paragraph("4.2 VEE Rules", style='Heading 2')
        new_doc.add_paragraph("4.2.1 INT-SPIKE-CHK", style='Heading 3')
        new_doc.add_paragraph("4.2.1.1 BO Algorithm")
        new_doc.add_paragraph("4.2.1.2 Before")
        new_doc.add_paragraph("4.2.1.3 After")
        new_doc.add_paragraph("4.2.1.4 DB Results")
        new_doc.add_paragraph("4.2.2 NEG-CONS-CHK", style='Heading 3')
        new_doc.add_paragraph("4.2.2.1 BO Algorithm")
        new_doc.add_paragraph("4.2.2.2 Before")
        new_doc.add_paragraph("4.2.2.3 After")
        new_doc.add_paragraph("4.2.2.4 DB Results")
        continue

    if within_execution:
        continue

    if copying:
        new_doc.add_paragraph(para.text, style=para.style)

# Save the modified document
output_path = "PGE_C2M_Unit Test Design_SOX_VEE_USG_V0.6_Corrected.docx"
new_doc.save(output_path)


